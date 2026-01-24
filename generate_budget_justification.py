#!/usr/bin/env python3
"""
NIH Budget Justification Generator for Northern Arizona University
Generates LaTeX and Microsoft Word (.docx) files from NAU budget Excel templates
Supports 3, 5, and 10-year project templates

Requirements:
    - openpyxl (Python package)
    - python-docx (Python package, for .docx formatting)
    - xelatex (for PDF compilation)
    - pandoc (optional, for .docx generation)

Usage:
    python3 generate_budget_justification.py MyBudget.xlsx
    python3 generate_budget_justification.py MyBudget.xlsx -o output_directory -v

Output:
    - MyBudget_BudgetJustification.tex (LaTeX source)
    - MyBudget_BudgetJustification.docx (Word document, if pandoc is installed)
"""

import openpyxl
import argparse
import os
import sys
import subprocess
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ============================================================================
# UNIVERSITY-MANDATED VERBIAGE CONSTANTS
# ============================================================================

FRINGE_BENEFITS_TEXT = """Fringe benefit rates are rounded estimates based on the projected cost of health, dental, life, disability, FICA and Medicare, unemployment, and retirement benefits relative to the employee's salary and/or wages, FTE, and election of benefits. The employee's fringe benefit rate is calculated by dividing their salary by the total cost of their benefits package."""

def get_indirect_costs_text(rate, mtdc_base, total_indirect):
    """Generate exact NAU MTDC indirect costs verbiage"""
    # Format currency with LaTeX escaping
    mtdc_formatted = format_currency(mtdc_base)
    indirect_formatted = format_currency(total_indirect)
    return f"""Indirect costs are calculated at {rate}\\% of the Modified Total Direct Cost (MTDC) base, per Northern Arizona University's federally negotiated rate agreement for on-campus research. The MTDC base of {mtdc_formatted} excludes equipment, participant support costs, tuition remission, and subaward amounts over \\$25,000. The total indirect costs for the project are {indirect_formatted}."""

# ============================================================================
# UTILITY FUNCTIONS
# ============================================================================

def format_currency(value):
    """Format number as currency with commas ($XX,XXX.XX or $XX,XXX)"""
    if value is None or value == '' or value == 0:
        return '\\$0'
    try:
        num = float(str(value).replace(',', ''))
        if num == int(num):
            return f'\\${int(num):,}'
        return f'\\${num:,.2f}'
    except:
        return '\\$0'

def format_year_range(year_details_list):
    """
    Format yearly breakdowns, using ranges for sequential years with same amounts.

    Args:
        year_details_list: List of tuples (year_num, amount)

    Returns:
        Formatted string like "Years 1-3: $1,000; Year 4: $2,000"
    """
    if not year_details_list:
        return ""

    # Group sequential years with same amount
    grouped = []
    i = 0
    while i < len(year_details_list):
        start_year, amount = year_details_list[i]
        end_year = start_year

        # Look ahead for sequential years with same amount
        j = i + 1
        while j < len(year_details_list):
            next_year, next_amount = year_details_list[j]
            if next_year == end_year + 1 and abs(next_amount - amount) < 0.01:  # Same amount
                end_year = next_year
                j += 1
            else:
                break

        # Format the range
        if start_year == end_year:
            grouped.append(f"Year {start_year}: {format_currency(amount)}")
        else:
            grouped.append(f"Years {start_year}-{end_year}: {format_currency(amount)}")

        i = j

    return '; '.join(grouped)

def format_docx_file(docx_path):
    """
    Format a .docx file to match PDF formatting:
    - 0.5 inch margins on all sides
    - Arial 10pt font
    - Justified text alignment
    - Red highlighting for TODO items

    Args:
        docx_path: Path to the .docx file to format
    """
    try:
        from docx.enum.text import WD_COLOR_INDEX
        from docx.shared import RGBColor

        doc = Document(docx_path)

        # Set margins to 0.5 inches on all sides
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)

        # Set font to Arial 10pt and justification for all paragraphs
        for paragraph in doc.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for run in paragraph.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)

                # Highlight TODO items in red
                if run.text and 'TODO' in run.text:
                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red text
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Yellow highlight

        # Format tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)

                            # Highlight TODO items in red
                            if run.text and 'TODO' in run.text:
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red text
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW  # Yellow highlight

        # Save the formatted document
        doc.save(docx_path)
        return True
    except Exception as e:
        return False

def escape_latex(text):
    """Escape LaTeX special characters"""
    if not isinstance(text, str):
        return str(text) if text is not None else ''
    replacements = {
        '#': '\\#',
        '$': '\\$',
        '%': '\\%',
        '&': '\\&',
        '_': '\\_',
        '{': '\\{',
        '}': '\\}',
        '~': '\\textasciitilde{}',
        '^': '\\textasciicircum{}'
    }
    result = text
    for char, replacement in replacements.items():
        result = result.replace(char, replacement)
    return result

def clean_numeric(value):
    """Convert Excel values to clean numbers"""
    if value is None or value == '':
        return 0
    try:
        return float(str(value).replace(',', ''))
    except:
        return 0

def highlight_todo(text):
    """Wrap TODO text in LaTeX highlighting using textcolor for robust compatibility"""
    return f"\\textcolor{{red}}{{{text}}}"

# ============================================================================
# BUDGET DATA EXTRACTOR
# ============================================================================

class BudgetExtractor:
    """Extract budget data from NAU Excel templates"""

    def __init__(self, filepath):
        self.filepath = filepath
        self.wb = openpyxl.load_workbook(filepath, data_only=True)
        self.years = self.detect_years()
        self.senior_personnel = []
        self.other_personnel = []
        self.domestic_travel = []
        self.international_travel = []
        self.cumulative_data = {}
        self.summary_sheet = self.wb['Summary_of_Personnel Costs ']  # Note the trailing space

    def detect_years(self):
        """Auto-detect 3, 5, or 10 year template by scanning year columns"""
        sheet = self.wb['Budget Details']
        # Scan Row 2, Columns 7-25 for year headers
        year_count = 0
        for col in range(7, 26):
            cell_value = sheet.cell(2, col).value
            if cell_value is not None:
                cell_str = str(cell_value).strip()
                # Check if it's a year label (Year 0, Year 1, etc.)
                if cell_str.startswith('Year') or cell_str.isdigit():
                    year_count += 1
        return year_count if year_count > 0 else 5  # Default to 5 if detection fails

    def extract_senior_personnel(self):
        """Extract all senior personnel dynamically from Row 11+"""
        sheet = self.wb['Budget Details']
        row = 11  # Fixed starting row for senior personnel

        while row < 50:  # Safety limit
            name = sheet.cell(row, 2).value
            role = sheet.cell(row, 4).value

            # Stop conditions: empty name, "Total" row, or "OTHER PERSONNEL"
            if not name or str(name).strip().lower() == 'none':
                row += 1
                continue
            if 'Total' in str(name) or 'OTHER' in str(name):
                break

            # Check if role is valid (not "None" string)
            if name and role and str(role).strip().lower() != 'none':
                person = {
                    'name': str(name).strip(),
                    'role': str(role).strip(),
                    'base_salary': clean_numeric(sheet.cell(row, 5).value),
                    'contract_type': clean_numeric(sheet.cell(row, 6).value),
                }

                # Extract person-months for each year (columns 7, 8, 9, 10, 11 for 5-year)
                for i in range(self.years):
                    pm_value = clean_numeric(sheet.cell(row, 7 + i).value)
                    person[f'pm_y{i+1}'] = pm_value

                # Extract salary for each year (columns 20, 21, 22, 23, 24 for 5-year)
                for i in range(self.years):
                    salary_value = clean_numeric(sheet.cell(row, 20 + i).value)
                    person[f'salary_y{i+1}'] = salary_value

                # Extract total salary (column 25)
                person['total_salary'] = clean_numeric(sheet.cell(row, 25).value)

                # Extract total compensation from Summary sheet to calculate fringe
                # Summary sheet starts at row 6 for first senior person (row 11 in Budget Details)
                summary_row = row - 11 + 6
                total_compensation = clean_numeric(self.summary_sheet.cell(summary_row, 9).value)  # Column I = total

                # Calculate fringe as total_compensation - total_salary
                if total_compensation and person['total_salary']:
                    person['total_fringe'] = total_compensation - person['total_salary']
                    # Calculate ERE rate
                    if person['total_salary'] > 0:
                        person['ere_rate'] = (person['total_fringe'] / person['total_salary']) * 100
                    else:
                        person['ere_rate'] = 0
                else:
                    person['total_fringe'] = 0
                    person['ere_rate'] = 0

                self.senior_personnel.append(person)

            row += 1

        return self.senior_personnel

    def extract_other_personnel(self):
        """Extract other personnel from Row 31-32+ area"""
        sheet = self.wb['Budget Details']

        # Find "OTHER PERSONNEL" header (scan rows 30-35)
        start_row = None
        for row in range(30, 36):
            num_cell = sheet.cell(row, 2).value
            if num_cell and 'OTHER' in str(num_cell):
                start_row = row + 2  # Data starts 2 rows after header (skip header row with #, Role, Hours)
                break

        if not start_row:
            return []

        row = start_row
        while row < start_row + 30:  # Max 30 other personnel rows
            number = sheet.cell(row, 2).value
            role = sheet.cell(row, 3).value

            # Stop if role is empty or "None"
            if not role or str(role).strip().lower() == 'none':
                row += 1
                continue

            # Must have a valid role description
            if role:
                # Check if number is valid (not "None")
                num_val = str(number).strip() if number else ''
                if num_val.lower() == 'none':
                    num_val = ''

                # Check if this is a GRA (monthly salary vs hourly)
                hours_cell = sheet.cell(row, 4).value
                is_gra = 'gra' in str(role).lower()

                # For GRAs, hours_week column may contain semester info (e.g., "Spring")
                # Rate is monthly salary, not hourly
                if is_gra and isinstance(hours_cell, str):
                    # GRA with semester designation - default to 20 hours/week
                    hours_week = 20
                    is_monthly = True
                    semester = hours_cell
                else:
                    hours_week = clean_numeric(hours_cell)
                    is_monthly = False
                    semester = None

                position = {
                    'number': num_val,
                    'role': str(role).strip(),
                    'hours_week': hours_week,
                    'rate': clean_numeric(sheet.cell(row, 5).value),
                    'is_monthly': is_monthly,
                    'semester': semester,
                }

                # For Other Personnel, ERE is not stored in columns 6-8 (those are months worked)
                # We'll calculate fringe from Summary sheet or Cumulative data later
                position['ere_rate'] = 0  # Placeholder, will be calculated from actual fringe

                # Extract salary for each year and total
                for i in range(self.years):
                    salary_value = clean_numeric(sheet.cell(row, 20 + i).value)
                    position[f'salary_y{i+1}'] = salary_value

                # Extract total salary (column 25)
                position['total_salary'] = clean_numeric(sheet.cell(row, 25).value)

                # Find matching row in Summary sheet to get fringe
                # Summary sheet has Other Personnel starting around row 31
                total_compensation = None
                for summary_row in range(30, 50):
                    summary_role = self.summary_sheet.cell(summary_row, 2).value
                    if summary_role and str(summary_role).strip() == str(role).strip():
                        total_compensation = clean_numeric(self.summary_sheet.cell(summary_row, 9).value)
                        break

                # Calculate fringe as total_compensation - total_salary
                if total_compensation and position['total_salary']:
                    position['total_fringe'] = total_compensation - position['total_salary']
                    # Calculate ERE rate
                    if position['total_salary'] > 0:
                        position['ere_rate'] = (position['total_fringe'] / position['total_salary']) * 100
                    else:
                        position['ere_rate'] = 0
                else:
                    # Fallback to blended rate if not found in Summary
                    cum_sheet = self.wb['Cumulative']
                    total_oth_salary = cum_sheet.cell(11, 36).value or 0
                    total_oth_benefits = cum_sheet.cell(14, 36).value or 0

                    if total_oth_salary > 0:
                        blended_ere = (total_oth_benefits / total_oth_salary) * 100
                        position['ere_rate'] = blended_ere
                        position['total_fringe'] = position['total_salary'] * (blended_ere / 100)
                    else:
                        position['ere_rate'] = 0
                        position['total_fringe'] = 0

                self.other_personnel.append(position)

            row += 1

        return self.other_personnel

    def extract_travel(self):
        """Extract travel data from Travel Calculator sheet, categorize using Budget Details"""
        # First, get travel categories from Budget Details sheet
        domestic_amounts = set()
        international_amounts = set()

        try:
            details_sheet = self.wb['Budget Details']
            current_category = None

            # Scan travel section in Budget Details (rows 158-170)
            for row in range(158, 175):
                cell_b = details_sheet.cell(row, 2).value
                cell_d = details_sheet.cell(row, 4).value
                year1_amt = clean_numeric(details_sheet.cell(row, 20).value)

                if cell_b:
                    cell_b_str = str(cell_b).strip().lower()
                    if 'domestic' in cell_b_str:
                        current_category = 'domestic'
                    elif 'international' in cell_b_str or 'foreign' in cell_b_str:
                        current_category = 'international'
                    elif 'total' in cell_b_str:
                        current_category = None

                # If we have a Year 1 amount and a current category, record it
                if current_category and year1_amt > 0:
                    if current_category == 'domestic':
                        domestic_amounts.add(year1_amt)
                    else:
                        international_amounts.add(year1_amt)
        except Exception:
            pass  # Fall back to heuristic if Budget Details parsing fails

        # Now extract trips from Travel Calculator
        try:
            sheet = self.wb[' Travel Calculator']  # Note: space before 'Travel'
        except KeyError:
            try:
                sheet = self.wb['Travel Calculator']  # Try without space
            except KeyError:
                return []

        # Rows 4-13 contain trip data in the template
        for row_idx in range(4, 14):
            destination = sheet.cell(row_idx, 2).value
            travelers = clean_numeric(sheet.cell(row_idx, 3).value)

            if destination and travelers and travelers > 0:
                trip = {
                    'destination': str(destination).strip(),
                    'travelers': int(travelers),
                    'days': clean_numeric(sheet.cell(row_idx, 4).value),
                    'nights': clean_numeric(sheet.cell(row_idx, 5).value),
                    'flight_pp': clean_numeric(sheet.cell(row_idx, 6).value),
                    'full_meal_pp': clean_numeric(sheet.cell(row_idx, 7).value),
                    'first_last_meal_pp': clean_numeric(sheet.cell(row_idx, 8).value),
                    'lodging_pn': clean_numeric(sheet.cell(row_idx, 9).value),
                    'transportation_pp': clean_numeric(sheet.cell(row_idx, 10).value),
                    'conf_reg_pp': clean_numeric(sheet.cell(row_idx, 11).value),
                    'misc_pp': clean_numeric(sheet.cell(row_idx, 12).value),  # Column 12 is miscellaneous per person
                    'total_first_last_meal': clean_numeric(sheet.cell(row_idx, 13).value),
                    'total_remaining_meal': clean_numeric(sheet.cell(row_idx, 14).value),
                    'total_flight': clean_numeric(sheet.cell(row_idx, 15).value),
                    'total_lodging': clean_numeric(sheet.cell(row_idx, 16).value),
                    'total_transportation': clean_numeric(sheet.cell(row_idx, 17).value),  # Column 17 is ground transportation total
                    'total_conf_reg': clean_numeric(sheet.cell(row_idx, 18).value),
                    'total_misc': clean_numeric(sheet.cell(row_idx, 19).value),
                    'cumulative': clean_numeric(sheet.cell(row_idx, 20).value),
                }

                # Categorize using Budget Details data
                # Check Year 1 amounts to match with Budget Details categories
                trip_year1 = clean_numeric(sheet.cell(row_idx, 20).value)  # Use cumulative as proxy

                # Match to international if amount appears in international set from Budget Details
                is_international = False
                if international_amounts:
                    is_international = trip_year1 in international_amounts

                # If no match found in international, check if explicitly in domestic
                if not is_international and domestic_amounts:
                    is_international = trip_year1 not in domestic_amounts and trip_year1 not in international_amounts
                    # If not found in either, fall back to heuristic
                    if trip_year1 not in domestic_amounts and trip_year1 not in international_amounts:
                        dest_lower = trip['destination'].lower()
                        is_international = 'international' in dest_lower or trip['flight_pp'] > 1000

                if is_international:
                    self.international_travel.append(trip)
                else:
                    self.domestic_travel.append(trip)

        return self.domestic_travel + self.international_travel

    def extract_cumulative(self):
        """Extract totals from Cumulative sheet (fixed row positions)"""
        sheet = self.wb['Cumulative']

        budget_items = {
            9: 'total_salaries',
            10: 'sr_personnel_salary',
            11: 'other_personnel_salary',
            12: 'total_benefits',
            13: 'sr_personnel_benefits',
            14: 'other_personnel_benefits',
            15: 'tuition_remission',
            16: 'total_equipment',
            17: 'total_travel',
            18: 'domestic_travel',
            19: 'international_travel',
            20: 'participant_support',
            21: 'other_direct_costs',
            22: 'materials_supplies',
            23: 'consultants',
            24: 'other_expenses',
            25: 'total_subawards',
            27: 'total_direct_costs',
            28: 'indirect_base',
            29: 'total_indirect_costs'
        }

        for row_idx, category in budget_items.items():
            years = {}
            # Total is in column 36
            total = clean_numeric(sheet.cell(row_idx, 36).value)

            # Years at columns 11, 16, 21, 26, 31, ... (+5 pattern)
            for i in range(self.years):
                year_col = 11 + (i * 5)
                years[f'year{i+1}'] = clean_numeric(sheet.cell(row_idx, year_col).value)

            self.cumulative_data[category] = {
                'total': total,
                **years
            }

        return self.cumulative_data

    def extract_other_direct_costs_items(self):
        """Extract individual Other Direct Costs line items from Budget Details rows 177-192"""
        self.odc_items = []  # Simple list of all ODC line items

        try:
            sheet = self.wb['Budget Details']
            current_category = None

            # Extract from specific rows 177-192 as requested
            for row in range(177, 193):
                cell_b = sheet.cell(row, 2).value
                cell_d = sheet.cell(row, 4).value

                # Check for category header in column B
                if cell_b:
                    cell_b_str = str(cell_b).strip()
                    if cell_b_str and cell_b_str.lower() not in ['none', '']:
                        current_category = cell_b_str

                # Get description from column D
                description = None
                if cell_d:
                    desc_str = str(cell_d).strip()
                    if desc_str and desc_str.lower() not in ['none', '', 'description']:
                        description = desc_str

                # Skip if no description
                if not description:
                    continue

                # Extract yearly amounts (columns 20-24)
                item = {
                    'category': current_category if current_category else '',
                    'description': description,
                    'yearly': {},
                    'total': 0
                }

                # Get yearly values
                for i in range(self.years):
                    year_val = clean_numeric(sheet.cell(row, 20 + i).value)
                    if year_val > 0:
                        item['yearly'][f'year{i+1}'] = year_val

                # Calculate total from sum of yearly amounts (column 25 contains category subtotals, not line item totals)
                item['total'] = sum(item['yearly'].values())

                # Only add if there's actual budget
                if item['total'] > 0:
                    self.odc_items.append(item)

        except Exception as e:
            pass  # Silently fail if extraction doesn't work

        return self.odc_items

    def extract_subaward_names(self):
        """Extract subaward organization names from Cumulative sheet"""
        self.subaward_names = []
        try:
            sheet = self.wb['Cumulative']
            # Look for subaward organization names in row 25 area or scan for text
            # Check rows 25-30 for subaward organization names in column 2
            for row in range(25, 35):
                cell_val = sheet.cell(row, 2).value
                if cell_val and isinstance(cell_val, str):
                    cell_str = cell_val.strip()
                    # Skip generic labels
                    if cell_str and 'subaward' not in cell_str.lower() and 'total' not in cell_str.lower():
                        # Check if this row has budget amounts (indicates it's an org name)
                        has_amount = False
                        for col in range(11, 40):
                            amt = clean_numeric(sheet.cell(row, col).value)
                            if amt > 0:
                                has_amount = True
                                break
                        if has_amount and len(cell_str) > 2:
                            self.subaward_names.append(cell_str)
        except Exception:
            pass
        return self.subaward_names

# ============================================================================
# LATEX GENERATOR
# ============================================================================

class LaTeXGenerator:
    """Generate LaTeX budget justification from extracted data"""

    def __init__(self, extractor):
        self.data = extractor

    def generate_header(self):
        """Generate document header"""
        return "\\chapter*{BUDGET JUSTIFICATION – NORTHERN ARIZONA UNIVERSITY}"

    def generate_personnel_section(self):
        """Section A: Senior/Key Personnel - fully dynamic"""
        total = self.data.cumulative_data.get('sr_personnel_salary', {}).get('total', 0)

        section = f"\\subsection*{{A. Senior Personnel—{format_currency(total)}}}\n"

        if not self.data.senior_personnel:
            section += "No senior personnel budgeted for this project.\n\n"
            return section

        for i, person in enumerate(self.data.senior_personnel, 1):
            name = escape_latex(person['name'])
            role = escape_latex(person['role'])
            pm_y1 = person.get('pm_y1', 0)
            person_total = person.get('total_salary', 0)
            base_salary = person.get('base_salary', 0)

            # Check if person months vary across years
            pm_values = [person.get(f'pm_y{j+1}', 0) for j in range(self.data.years)]
            pm_all_same = len(set(pm_values)) <= 1  # All values are the same

            # Build the person months description for the header
            if pm_all_same:
                pm_description = f"{pm_y1} Months per year"
            else:
                # Show variation with ranges: "Person months: 12 in Year 1, 3 in Years 2-5"
                pm_year_data = [(j+1, person.get(f'pm_y{j+1}', 0)) for j in range(self.data.years) if person.get(f'pm_y{j+1}', 0) > 0]
                # Format as "X months in Year/Years Y"
                grouped = []
                i = 0
                while i < len(pm_year_data):
                    start_year, months = pm_year_data[i]
                    end_year = start_year

                    # Look ahead for sequential years with same months
                    j = i + 1
                    while j < len(pm_year_data):
                        next_year, next_months = pm_year_data[j]
                        if next_year == end_year + 1 and abs(next_months - months) < 0.01:
                            end_year = next_year
                            j += 1
                        else:
                            break

                    # Format the range
                    if start_year == end_year:
                        grouped.append(f"{months} in Year {start_year}")
                    else:
                        grouped.append(f"{months} in Years {start_year}-{end_year}")

                    i = j

                pm_description = f"Person months: {', '.join(grouped)}"

            section += f"\\subsubsection*{{A{i}. {name}, {role}: {pm_description}, {format_currency(person_total)} total}}\n"

            # Generate role-specific narrative (starts on new line after header)
            last_name = name.split()[-1]

            if 'PI' in role and 'Co' not in role:
                # Principal Investigator - detailed narrative
                section += f"{last_name} is the Principal Investigator who will provide overall leadership, direction, and coordination for all aspects of this research project. "
                section += f"{last_name} will be responsible for scientific and administrative oversight, ensuring that project milestones are met, coordinating with collaborators, managing the research team, and ensuring compliance with all institutional and funding agency requirements. "

                # Describe effort commitment based on whether it varies
                if pm_all_same:
                    section += f"{last_name} will dedicate {pm_y1} person months per year to this project, with a base salary of {format_currency(base_salary)}. "
                else:
                    # Show effort by year
                    effort_parts = [f"{person.get(f'pm_y{j+1}', 0)} person months in Year {j+1}" for j in range(self.data.years) if person.get(f'pm_y{j+1}', 0) > 0]
                    section += f"{last_name} will dedicate {', '.join(effort_parts[:-1])}, and {effort_parts[-1]} to this project, with a base salary of {format_currency(base_salary)}. " if len(effort_parts) > 1 else f"{last_name} will dedicate {effort_parts[0]} to this project, with a base salary of {format_currency(base_salary)}. "


                # Add year-by-year breakdown if available
                year_data = []
                for i in range(self.data.years):
                    salary = person.get(f'salary_y{i+1}', 0)
                    if salary > 0:
                        year_data.append((i+1, salary))
                if year_data:
                    section += f"The total salary requested for {last_name} over {self.data.years} years is {format_currency(person_total)}, allocated as follows: {format_year_range(year_data)}. "
                    section += "A 3\\% annual salary increase is included. "

                section += f"{last_name}'s expertise and leadership are essential to the success of this project.\n\n"

            elif 'Co' in role and 'PI' in role:
                # Co-PI - detailed narrative
                section += f"{last_name} serves as Co-Principal Investigator and will play a critical role in the scientific direction and execution of this research. "

                # Describe effort commitment based on whether it varies
                if pm_all_same:
                    section += f"{last_name} will contribute {pm_y1} person months per year to the project, bringing essential expertise and working in close collaboration with the PI to ensure project success. "
                else:
                    # Show effort by year
                    effort_parts = [f"{person.get(f'pm_y{j+1}', 0)} person months in Year {j+1}" for j in range(self.data.years) if person.get(f'pm_y{j+1}', 0) > 0]
                    section += f"{last_name} will contribute {', '.join(effort_parts[:-1])}, and {effort_parts[-1]} to the project, bringing essential expertise and working in close collaboration with the PI to ensure project success. " if len(effort_parts) > 1 else f"{last_name} will contribute {effort_parts[0]} to the project, bringing essential expertise and working in close collaboration with the PI to ensure project success. "

                section += f"With a base salary of {format_currency(base_salary)}, {last_name} will share responsibility for key project decisions, mentor junior team members, contribute to data analysis and interpretation, and assist in the preparation of publications and presentations. "

                # Add year-by-year breakdown
                year_data = []
                for i in range(self.data.years):
                    salary = person.get(f'salary_y{i+1}', 0)
                    if salary > 0:
                        year_data.append((i+1, salary))
                if year_data:
                    section += f"The total salary requested for {last_name} is {format_currency(person_total)}, distributed as: {format_year_range(year_data)}. "
                    section += "A 3\\% annual salary increase is included. "

                section += highlight_todo(f"[TODO: Describe {last_name}'s specific technical expertise, prior relevant experience, and unique contributions to this project.]") + "\n\n"

            else:
                # Senior Personnel - detailed narrative
                section += f"{last_name} will contribute critical expertise to this research project as Senior Personnel. "

                # Describe effort commitment based on whether it varies
                if pm_all_same:
                    section += f"{last_name} is requesting {pm_y1} months of support per year, with a base salary of {format_currency(base_salary)}. "
                else:
                    # Show effort by year
                    effort_parts = [f"{person.get(f'pm_y{j+1}', 0)} months in Year {j+1}" for j in range(self.data.years) if person.get(f'pm_y{j+1}', 0) > 0]
                    section += f"{last_name} is requesting {', '.join(effort_parts[:-1])}, and {effort_parts[-1]} of support, with a base salary of {format_currency(base_salary)}. " if len(effort_parts) > 1 else f"{last_name} is requesting {effort_parts[0]} of support, with a base salary of {format_currency(base_salary)}. "

                section += f"In this role, {last_name} will provide specialized knowledge and technical guidance, assist with specific research tasks, participate in project meetings and strategic planning, and contribute to the dissemination of research findings. "

                # Add year-by-year breakdown
                year_data = []
                for i in range(self.data.years):
                    salary = person.get(f'salary_y{i+1}', 0)
                    if salary > 0:
                        year_data.append((i+1, salary))
                if year_data:
                    section += f"The total salary requested for {last_name} is {format_currency(person_total)}, with yearly allocation: {format_year_range(year_data)}. "
                    section += "A 3\\% annual salary increase is included. "

                section += highlight_todo(f"[TODO: Describe {last_name}'s specific role, specialized qualifications, and how their expertise complements the research team.]") + "\n\n"

        return section

    def _get_years_intro(self, years_with_salary, hours, rate, position_type="position"):
        """Helper to generate consistent year-aware introduction for Other Personnel"""
        all_years_active = len(years_with_salary) == self.data.years

        if all_years_active:
            return f"will be hired to work {hours} hours per week at a rate of {format_currency(rate)} per hour"

        if len(years_with_salary) == 1:
            return f"will be hired in Year {years_with_salary[0]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour"
        elif len(years_with_salary) > 2 and years_with_salary == list(range(years_with_salary[0], years_with_salary[-1] + 1)):
            return f"will be hired in Years {years_with_salary[0]}-{years_with_salary[-1]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour"
        else:
            years_str = ', '.join(str(y) for y in years_with_salary)
            return f"will be hired in Years {years_str} to work {hours} hours per week at a rate of {format_currency(rate)} per hour"

    def generate_other_personnel_section(self):
        """Section B: Other Personnel"""
        total_salary = self.data.cumulative_data.get('other_personnel_salary', {}).get('total', 0)
        total_benefits = self.data.cumulative_data.get('other_personnel_benefits', {}).get('total', 0)
        tuition = self.data.cumulative_data.get('tuition_remission', {}).get('total', 0)

        # Calculate grand total
        grand_total = total_salary + total_benefits + tuition

        section = f"\\subsection*{{B. Other Personnel—{format_currency(grand_total)}}}\n"

        # Filter to only include personnel with budget amounts > 0
        budgeted_personnel = [p for p in self.data.other_personnel if p.get('total_salary', 0) > 0]

        if not budgeted_personnel:
            section += "No other personnel budgeted for this project.\n\n"
            return section

        for i, position in enumerate(budgeted_personnel, 1):
            role = escape_latex(position['role'])
            hours = position.get('hours_week', 0)
            position_total = position.get('total_salary', 0)
            rate = position.get('rate', 0)

            # Check which years have salary (indicating when position is active)
            years_with_salary = [j+1 for j in range(self.data.years) if position.get(f'salary_y{j+1}', 0) > 0]
            all_years_active = len(years_with_salary) == self.data.years

            # Build the effort description for the header
            if all_years_active:
                effort_description = f"{hours} hours/week"
            else:
                # Show which years: "20 hours/week (Years 1-2)" or "20 hours/week (Years 1, 3, 5)"
                if len(years_with_salary) > 2 and years_with_salary == list(range(years_with_salary[0], years_with_salary[-1] + 1)):
                    # Consecutive years
                    effort_description = f"{hours} hours/week (Years {years_with_salary[0]}-{years_with_salary[-1]})"
                else:
                    # Non-consecutive years
                    years_str = ', '.join(str(y) for y in years_with_salary)
                    effort_description = f"{hours} hours/week (Years {years_str})"

            section += f"\\subsubsection*{{B{i}. {role}: {effort_description}, {format_currency(position_total)} total}}\n"

            # Generate description based on role type with detailed narratives
            role_lower = role.lower()

            # Check for undergrad BEFORE checking for grad (to avoid matching "undergrad" as "grad")
            if 'undergrad' in role_lower:
                # Add time frame if not all years
                if all_years_active:
                    section += f"Undergraduate students will be hired to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                else:
                    if len(years_with_salary) == 1:
                        section += f"Undergraduate students will be hired in Year {years_with_salary[0]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                    elif len(years_with_salary) > 2 and years_with_salary == list(range(years_with_salary[0], years_with_salary[-1] + 1)):
                        section += f"Undergraduate students will be hired in Years {years_with_salary[0]}-{years_with_salary[-1]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                    else:
                        years_str = ', '.join(str(y) for y in years_with_salary)
                        section += f"Undergraduate students will be hired in Years {years_str} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "

                section += f"Undergraduate researchers will gain hands-on research experience by assisting with data collection and entry, conducting literature searches, preparing research materials, maintaining laboratory notebooks, participating in team meetings, and contributing to specific project tasks under the supervision of senior personnel. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total undergraduate support is budgeted at {format_currency(position_total)}, with yearly allocation: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "This investment in undergraduate training aligns with the university's educational mission and provides essential research support. "
                section += highlight_todo("[TODO: Specify training objectives and anticipated student contributions]") + "\n\n"

            elif 'postdoc' in role_lower or 'post-doc' in role_lower:
                # Add time frame if not all years
                if all_years_active:
                    section += f"A postdoctoral researcher will be hired to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                else:
                    if len(years_with_salary) == 1:
                        section += f"A postdoctoral researcher will be hired in Year {years_with_salary[0]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                    elif len(years_with_salary) > 2 and years_with_salary == list(range(years_with_salary[0], years_with_salary[-1] + 1)):
                        section += f"A postdoctoral researcher will be hired in Years {years_with_salary[0]}-{years_with_salary[-1]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                    else:
                        years_str = ', '.join(str(y) for y in years_with_salary)
                        section += f"A postdoctoral researcher will be hired in Years {years_str} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "

                section += f"The postdoctoral researcher will conduct independent research under the guidance of the PI, assist with experimental design and data collection, train and supervise graduate and undergraduate students, contribute to manuscript preparation and grant writing, and participate in lab meetings and professional development activities. "

                # Add year-by-year breakdown
                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"The total salary requested for this position is {format_currency(position_total)}, distributed as: {format_year_range(year_data)}. "

                section += "This position is essential for maintaining research productivity and providing mentorship to junior team members. "
                section += highlight_todo("[TODO: Specify research focus areas and required qualifications]") + "\n\n"

            elif 'grad' in role_lower or 'gra' in role_lower:
                # Check if this is monthly or hourly pay
                is_monthly = position.get('is_monthly', False)
                rate_text = f"a monthly stipend of {format_currency(rate)}" if is_monthly else f"a rate of {format_currency(rate)} per hour"

                # Add time frame if not all years
                if all_years_active:
                    section += f"A Graduate Research Assistant will be hired to work {hours} hours per week at {rate_text}. "
                else:
                    if len(years_with_salary) == 1:
                        section += f"A Graduate Research Assistant will be hired in Year {years_with_salary[0]} to work {hours} hours per week at {rate_text}. "
                    elif len(years_with_salary) > 2 and years_with_salary == list(range(years_with_salary[0], years_with_salary[-1] + 1)):
                        section += f"A Graduate Research Assistant will be hired in Years {years_with_salary[0]}-{years_with_salary[-1]} to work {hours} hours per week at {rate_text}. "
                    else:
                        years_str = ', '.join(str(y) for y in years_with_salary)
                        section += f"A Graduate Research Assistant will be hired in Years {years_str} to work {hours} hours per week at {rate_text}. "

                section += f"The GRA will assist with literature reviews, data collection and analysis, laboratory experiments, field work as needed, maintenance of research equipment and supplies, preparation of research presentations, and participation in project team meetings. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total support for this position is {format_currency(position_total)}, allocated as: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "This position will provide valuable research training and contribute significantly to project deliverables. "
                section += highlight_todo("[TODO: Specify required academic background and specific responsibilities]") + "\n\n"

            elif 'project' in role_lower and 'manager' in role_lower:
                intro = self._get_years_intro(years_with_salary, hours, rate)
                section += f"A Project Manager {intro}. "
                section += f"The Project Manager will coordinate project activities across team members and collaborating institutions, manage the project timeline and deliverables, organize team meetings and communications, oversee budget expenditures and financial reporting, ensure compliance with institutional and funding agency requirements, maintain project documentation and databases, and facilitate dissemination of research findings. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total compensation for this position is {format_currency(position_total)}, distributed over {self.data.years} years as: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "This position is critical for ensuring efficient project execution and successful completion of all milestones. "
                section += highlight_todo("[TODO: Specify required project management experience and qualifications]") + "\n\n"

            elif 'lab' in role_lower and 'manager' in role_lower:
                intro = self._get_years_intro(years_with_salary, hours, rate)
                section += f"A Laboratory Manager {intro}. "
                section += f"The Lab Manager will maintain laboratory equipment and facilities, manage laboratory supplies and inventory, ensure compliance with safety regulations and protocols, train personnel on equipment use and safety procedures, coordinate equipment maintenance and repairs, maintain laboratory records and documentation, and support day-to-day laboratory operations. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total funding for this position is {format_currency(position_total)}, allocated as: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "This position is essential for maintaining a safe, efficient, and productive laboratory environment. "
                section += highlight_todo("[TODO: Specify required technical skills and laboratory experience]") + "\n\n"

            elif 'engineer' in role_lower or 'research' in role_lower and 'specialist' in role_lower:
                intro = self._get_years_intro(years_with_salary, hours, rate)
                section += f"A Research Engineer/Specialist {intro}. "
                section += f"This technical specialist will design and implement experimental protocols, develop and maintain research instrumentation and equipment, perform complex technical analyses, troubleshoot technical issues, provide technical training to research team members, contribute to method development and optimization, and assist with technical aspects of manuscript preparation. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total support is budgeted at {format_currency(position_total)}, with distribution: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "This specialized technical expertise is crucial for achieving the project's technical objectives. "
                section += highlight_todo("[TODO: Specify required technical specializations and experience]") + "\n\n"

            elif 'program' in role_lower and 'evaluator' in role_lower:
                intro = self._get_years_intro(years_with_salary, hours, rate)
                section += f"A Program Evaluator {intro}. "
                section += f"The Program Evaluator will design and implement evaluation frameworks and methodologies, collect and analyze program outcome data, assess project effectiveness and impact, prepare evaluation reports and recommendations, conduct stakeholder surveys and interviews, monitor progress toward project goals, and provide feedback to improve program implementation. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total evaluation support is {format_currency(position_total)}, allocated as: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "Independent evaluation is essential for assessing project impact and informing continuous improvement. "
                section += highlight_todo("[TODO: Specify evaluation methodologies and expected deliverables]") + "\n\n"

            elif 'data' in role_lower and ('analyst' in role_lower or 'scientist' in role_lower):
                intro = self._get_years_intro(years_with_salary, hours, rate)
                section += f"A Data Analyst/Scientist {intro}. "
                section += f"This position will manage and analyze research data, develop and implement data management protocols, perform statistical analyses and modeling, create data visualizations and reports, ensure data quality and integrity, maintain research databases, and contribute to data-related sections of publications and presentations. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total compensation is {format_currency(position_total)}, distributed as: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "Dedicated data analysis support is critical for extracting meaningful insights from research data. "
                section += highlight_todo("[TODO: Specify required analytical skills and software proficiency]") + "\n\n"

            elif 'student' in role_lower or 'worker' in role_lower:
                intro = self._get_years_intro(years_with_salary, hours, rate)
                section += f"Student workers {intro}. "
                section += f"These students will provide general research support including data entry and organization, literature searches and bibliographic management, preparation of research materials and supplies, assistance with routine laboratory or field tasks, participation in project activities under supervision, and other duties as assigned. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total student worker support is {format_currency(position_total)}, with allocation: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += "Student workers provide flexible research support while gaining valuable professional experience. "
                section += highlight_todo("[TODO: Specify anticipated tasks and training opportunities]") + "\n\n"

            elif 'temp' in role_lower or 'opt' in role_lower:
                intro = self._get_years_intro(years_with_salary, hours, rate)
                section += f"A temporary employee {intro}. "
                section += f"This position will provide flexible staffing support for specific project needs, assist with time-limited project tasks, cover staffing gaps as needed, and support project activities during peak work periods. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total budgeted support is {format_currency(position_total)}, allocated as: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += highlight_todo("[TODO: Describe specific responsibilities and support activities]") + "\n\n"

            else:
                # Generic other personnel - add time frame if not all years
                if all_years_active:
                    section += f"This {role} position will work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                else:
                    if len(years_with_salary) == 1:
                        section += f"This {role} position will be hired in Year {years_with_salary[0]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                    elif len(years_with_salary) > 2 and years_with_salary == list(range(years_with_salary[0], years_with_salary[-1] + 1)):
                        section += f"This {role} position will be hired in Years {years_with_salary[0]}-{years_with_salary[-1]} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "
                    else:
                        years_str = ', '.join(str(y) for y in years_with_salary)
                        section += f"This {role} position will be hired in Years {years_str} to work {hours} hours per week at a rate of {format_currency(rate)} per hour. "

                section += f"The position will support project activities, contribute to research objectives, participate in team meetings and collaborations, and perform assigned duties under the supervision of senior project personnel. "

                year_data = []
                for j in range(self.data.years):
                    salary = position.get(f'salary_y{j+1}', 0)
                    if salary > 0:
                        year_data.append((j+1, salary))
                if year_data:
                    section += f"Total salary for this position is {format_currency(position_total)}, distributed as: {format_year_range(year_data)}. "

                section += r"A 3\% annual salary increase is included. "

                section += highlight_todo("[TODO: Describe specific responsibilities, required qualifications, and contributions to project objectives]") + "\n\n"

        return section

    def generate_fringe_section(self):
        """Section C: Fringe Benefits with exact NAU verbiage - single paragraph format"""
        total = self.data.cumulative_data.get('total_benefits', {}).get('total', 0)

        section = f"\\subsection*{{C. Fringe Benefits—{format_currency(total)}}}\n"
        section += FRINGE_BENEFITS_TEXT + " "

        # Build fringe benefits as single paragraph - collect all entries
        fringe_entries = []

        # List fringe benefits for each senior personnel individual using actual ERE rates
        for person in self.data.senior_personnel:
            name = escape_latex(person['name'])
            last_name = name.split()[-1]
            person_salary = person.get('total_salary', 0)
            person_fringe = person.get('total_fringe', 0)
            ere_rate = person.get('ere_rate', 0)

            if person_salary > 0:
                # Use actual fringe from budget if available, otherwise calculate
                if person_fringe > 0:
                    # Calculate actual rate from budget values
                    actual_rate = (person_fringe / person_salary) * 100 if person_salary > 0 else ere_rate
                    fringe_entries.append(f"{last_name}'s fringe benefits are calculated at {actual_rate:.1f}\\%, totaling {format_currency(person_fringe)}")
                elif ere_rate > 0:
                    calculated_fringe = person_salary * (ere_rate / 100)
                    fringe_entries.append(f"{last_name}'s fringe benefits are calculated at {ere_rate:.1f}\\%, totaling {format_currency(calculated_fringe)}")

        # List fringe benefits for each other personnel position with budget
        for position in self.data.other_personnel:
            position_salary = position.get('total_salary', 0)
            position_fringe = position.get('total_fringe', 0)
            ere_rate = position.get('ere_rate', 0)

            if position_salary > 0:
                role = escape_latex(position['role'])
                # Use actual fringe from budget if available
                if position_fringe > 0:
                    actual_rate = (position_fringe / position_salary) * 100 if position_salary > 0 else ere_rate
                    fringe_entries.append(f"{role} fringe benefits are calculated at {actual_rate:.1f}\\%, totaling {format_currency(position_fringe)}")
                elif ere_rate > 0:
                    calculated_fringe = position_salary * (ere_rate / 100)
                    fringe_entries.append(f"{role} fringe benefits are calculated at {ere_rate:.1f}\\%, totaling {format_currency(calculated_fringe)}")

        # Join all entries with periods and spaces (single paragraph)
        if fringe_entries:
            section += ". ".join(fringe_entries) + ". "

        section += f"The total fringe benefits for all personnel over {self.data.years} years are {format_currency(total)}.\n\n"

        return section

    def generate_equipment_section(self):
        """Section D: Equipment"""
        total = self.data.cumulative_data.get('total_equipment', {}).get('total', 0)

        section = f"\\subsection*{{D. Equipment—"
        if total == 0:
            section += "N/A}\n"
            section += "No equipment over \\$5,000 is requested for this project.\n\n"
        else:
            section += f"{format_currency(total)}}}\n"
            section += highlight_todo("[TODO: List equipment items over \\$5,000 with justifications]") + "\n\n"

        return section

    def generate_domestic_travel_table(self):
        """Generate table for domestic trips only"""
        if not self.data.domestic_travel:
            return ""

        trips = self.data.domestic_travel[:5]  # Limit to 5 trips for table width
        num_trips = len(trips)

        if num_trips == 0:
            return ""

        # Build column specification
        col_spec = "|l|l|" + "r|" * num_trips

        table = f"\\begin{{table}}[h]\n\\centering\n\\caption{{Domestic Travel Budget Details}}\n\\label{{tab:domestic_travel}}\n\\small\n\\begin{{tabular}}{{{col_spec}}}\n\\hline\n"
        table += "\\rowcolor{gray!20}\n & \\textbf{Destination}"

        for trip in trips:
            table += f" & \\textbf{{{escape_latex(trip['destination'])}}}"
        table += " \\\\\n\\hline\n"

        # Info section (without Years row - yearly data is in the narrative)
        info_data = [
            ('Travelers', 'travelers'),
            ('Days', 'days'),
            ('Nights', 'nights')
        ]
        info_row_count = len(info_data)

        for idx, (label, key) in enumerate(info_data):
            if idx == 0:
                table += f"\\multirow{{{info_row_count}}}{{*}}{{\\textbf{{Info}}}} & {label}"
            else:
                table += f" & {label}"

            for trip in trips:
                val = trip.get(key, 0)
                table += f" & {int(val) if val else 0}"
            table += " \\\\\n"

        table += "\\hline\n"

        # Cost section - only include rows with at least one non-zero value
        cost_data = [
            ('Flight (per person)', 'flight_pp'),
            ('Full day meals (per person)', 'full_meal_pp'),
            ('1st/last day meals (per person)', 'first_last_meal_pp'),
            ('Accommodation (per night)', 'lodging_pn'),
            ('Conference reg. (per person)', 'conf_reg_pp'),
            ('Transportation (per person)', 'transportation_pp'),
            ('Miscellaneous (per person)', 'misc_pp')
        ]

        # Filter to only rows with data
        cost_rows_with_data = []
        for label, key in cost_data:
            if any(trip.get(key, 0) > 0 for trip in trips):
                cost_rows_with_data.append((label, key))

        if cost_rows_with_data:
            cost_row_count = len(cost_rows_with_data)
            for idx, (label, key) in enumerate(cost_rows_with_data):
                if idx == 0:
                    table += f"\\multirow{{{cost_row_count}}}{{*}}{{\\textbf{{Cost}}}} & {label}"
                else:
                    table += f" & {label}"

                for trip in trips:
                    val = trip.get(key, 0)
                    if val and val > 0:
                        table += f" & {format_currency(val)}"
                    else:
                        table += " & --"
                table += " \\\\\n"

            table += "\\hline\n"

        # Totals section - only include rows with at least one non-zero value
        totals_data = [
            ('1st/last day meal', 'total_first_last_meal'),
            ('Remaining days meal', 'total_remaining_meal'),
            ('Flight', 'total_flight'),
            ('Lodging', 'total_lodging'),
            ('Conference Registration', 'total_conf_reg'),
            ('Ground Transportation', 'total_transportation')
        ]

        # Filter to only rows with data
        totals_rows_with_data = []
        for label, key in totals_data:
            if any(trip.get(key, 0) > 0 for trip in trips):
                totals_rows_with_data.append((label, key))

        if totals_rows_with_data:
            totals_row_count = len(totals_rows_with_data)
            for idx, (label, key) in enumerate(totals_rows_with_data):
                if idx == 0:
                    table += f"\\multirow{{{totals_row_count}}}{{*}}{{\\textbf{{Totals}}}} & {label}"
                else:
                    table += f" & {label}"

                for trip in trips:
                    val = trip.get(key, 0)
                    if val and val > 0:
                        table += f" & {format_currency(val)}"
                    else:
                        table += " & --"
                table += " \\\\\n"

            table += "\\hline\n"

        # Cumulative row
        table += "\\textbf{Cumulative} & \\textbf{Cumulative}"
        for trip in trips:
            table += f" & \\textbf{{{format_currency(trip['cumulative'])}}}"
        table += " \\\\\n"

        table += "\\hline\n\\end{tabular}\n\\end{table}\n"

        return table

    def generate_international_travel_table(self):
        """Generate table for international trips only"""
        if not self.data.international_travel:
            return ""

        trips = self.data.international_travel[:5]  # Limit to 5 trips
        num_trips = len(trips)

        if num_trips == 0:
            return ""

        # Build column specification
        col_spec = "|l|l|" + "r|" * num_trips

        table = f"\\begin{{table}}[h]\n\\centering\n\\caption{{International Travel Budget Details}}\n\\label{{tab:international_travel}}\n\\small\n\\begin{{tabular}}{{{col_spec}}}\n\\hline\n"
        table += "\\rowcolor{gray!20}\n & \\textbf{Destination}"

        for trip in trips:
            table += f" & \\textbf{{{escape_latex(trip['destination'])}}}"
        table += " \\\\\n\\hline\n"

        # Info section (without Years row - yearly data is in the narrative)
        info_data = [
            ('Travelers', 'travelers'),
            ('Days', 'days'),
            ('Nights', 'nights')
        ]
        info_row_count = len(info_data)

        for idx, (label, key) in enumerate(info_data):
            if idx == 0:
                table += f"\\multirow{{{info_row_count}}}{{*}}{{\\textbf{{Info}}}} & {label}"
            else:
                table += f" & {label}"

            for trip in trips:
                val = trip.get(key, 0)
                table += f" & {int(val) if val else 0}"
            table += " \\\\\n"

        table += "\\hline\n"

        # Cost section - only include rows with at least one non-zero value
        cost_data = [
            ('Flight (per person)', 'flight_pp'),
            ('Full day meals (per person)', 'full_meal_pp'),
            ('1st/last day meals (per person)', 'first_last_meal_pp'),
            ('Accommodation (per night)', 'lodging_pn'),
            ('Conference reg. (per person)', 'conf_reg_pp'),
            ('Transportation (per person)', 'transportation_pp'),
            ('Miscellaneous (per person)', 'misc_pp')
        ]

        # Filter to only rows with data
        cost_rows_with_data = []
        for label, key in cost_data:
            if any(trip.get(key, 0) > 0 for trip in trips):
                cost_rows_with_data.append((label, key))

        if cost_rows_with_data:
            cost_row_count = len(cost_rows_with_data)
            for idx, (label, key) in enumerate(cost_rows_with_data):
                if idx == 0:
                    table += f"\\multirow{{{cost_row_count}}}{{*}}{{\\textbf{{Cost}}}} & {label}"
                else:
                    table += f" & {label}"

                for trip in trips:
                    val = trip.get(key, 0)
                    if val and val > 0:
                        table += f" & {format_currency(val)}"
                    else:
                        table += " & --"
                table += " \\\\\n"

            table += "\\hline\n"

        # Totals section - only include rows with at least one non-zero value
        totals_data = [
            ('1st/last day meal', 'total_first_last_meal'),
            ('Remaining days meal', 'total_remaining_meal'),
            ('Flight', 'total_flight'),
            ('Lodging', 'total_lodging'),
            ('Conference Registration', 'total_conf_reg'),
            ('Ground Transportation', 'total_transportation')
        ]

        # Filter to only rows with data
        totals_rows_with_data = []
        for label, key in totals_data:
            if any(trip.get(key, 0) > 0 for trip in trips):
                totals_rows_with_data.append((label, key))

        if totals_rows_with_data:
            totals_row_count = len(totals_rows_with_data)
            for idx, (label, key) in enumerate(totals_rows_with_data):
                if idx == 0:
                    table += f"\\multirow{{{totals_row_count}}}{{*}}{{\\textbf{{Totals}}}} & {label}"
                else:
                    table += f" & {label}"

                for trip in trips:
                    val = trip.get(key, 0)
                    if val and val > 0:
                        table += f" & {format_currency(val)}"
                    else:
                        table += " & --"
                table += " \\\\\n"

            table += "\\hline\n"

        # Cumulative row
        table += "\\textbf{Cumulative} & \\textbf{Cumulative}"
        for trip in trips:
            table += f" & \\textbf{{{format_currency(trip['cumulative'])}}}"
        table += " \\\\\n"

        table += "\\hline\n\\end{tabular}\n\\end{table}\n"

        return table

    def get_domestic_trip_narrative(self, dest, travelers, days, nights, trip, total_cost, trip_idx):
        """Generate narrative for domestic travel - conference if registration, otherwise site visit/field work options"""
        dest_escaped = escape_latex(dest)

        # Build cost breakdown string
        costs = []
        if trip['flight_pp'] > 0:
            costs.append(f"airfare ({format_currency(trip['total_flight'])} total)")
        if trip['total_lodging'] > 0:
            costs.append(f"lodging ({format_currency(trip['total_lodging'])})")
        if trip['total_first_last_meal'] + trip['total_remaining_meal'] > 0:
            meal_total = trip['total_first_last_meal'] + trip['total_remaining_meal']
            costs.append(f"meals ({format_currency(meal_total)})")
        if trip['total_conf_reg'] > 0:
            costs.append(f"conference registration ({format_currency(trip['total_conf_reg'])})")
        if trip['total_transportation'] > 0:
            costs.append(f"ground transportation ({format_currency(trip['total_transportation'])})")

        cost_str = ""
        if costs:
            if len(costs) == 1:
                cost_str = f"Costs include {costs[0]}. "
            else:
                cost_str = f"Costs include {', '.join(costs[:-1])} and {costs[-1]}. "

        # Check if conference registration is included
        has_conf_reg = trip.get('total_conf_reg', 0) > 0 or trip.get('conf_reg_pp', 0) > 0

        if has_conf_reg:
            # Conference travel - auto-detected from registration fee
            narrative = f"This trip to {dest_escaped} will support {travelers} traveler(s) attending a professional conference for {days} day(s) and {nights} night(s). "
            narrative += f"{cost_str}"
            narrative += f"The total cost for this conference trip is {format_currency(total_cost)} (see Table~\\ref{{tab:domestic_travel}}, Trip {trip_idx}). "
            narrative += "Conference attendance is essential for disseminating research findings, networking with peers, staying current with developments in the field, and receiving feedback on ongoing work. "
            narrative += highlight_todo("[TODO: Specify conference name, e.g., 'ACM CHI Conference on Human Factors in Computing Systems']")
        else:
            # No conference registration - show site visit AND field work options only
            narrative = f"This trip to {dest_escaped} will support {travelers} traveler(s) for {days} day(s) and {nights} night(s). "
            narrative += f"{cost_str}"
            narrative += f"The total cost for this trip is {format_currency(total_cost)} (see Table~\\ref{{tab:domestic_travel}}, Trip {trip_idx}).\n\n"

            # Field work justification
            narrative += highlight_todo("[IF FIELD WORK: Delete site visit justification below] ")
            narrative += "Field work is necessary for data collection, participant recruitment, conducting experiments in naturalistic settings, and gathering essential research materials that cannot be obtained remotely. "
            narrative += highlight_todo("[TODO: Specify field site location and research activities to be conducted]")
            narrative += "\n\n"

            # Site visit justification
            narrative += highlight_todo("[IF SITE VISIT: Delete field work justification above] ")
            narrative += "Site visits are critical for coordinating with collaborators, reviewing progress, troubleshooting technical issues, training personnel, and ensuring alignment of research activities across institutions. "
            narrative += highlight_todo("[TODO: Specify collaborating institution or partner organization]")

        return narrative

    def get_international_trip_narrative(self, dest, travelers, days, nights, trip, total_cost, trip_idx):
        """Generate narrative for international travel - always assume conference"""
        dest_escaped = escape_latex(dest)

        # Build cost breakdown string
        costs = []
        if trip['flight_pp'] > 0:
            costs.append(f"international airfare ({format_currency(trip['total_flight'])} total)")
        if trip['total_lodging'] > 0:
            costs.append(f"lodging ({format_currency(trip['total_lodging'])})")
        if trip['total_first_last_meal'] + trip['total_remaining_meal'] > 0:
            meal_total = trip['total_first_last_meal'] + trip['total_remaining_meal']
            costs.append(f"meals and per diem ({format_currency(meal_total)})")
        if trip['total_conf_reg'] > 0:
            costs.append(f"conference registration ({format_currency(trip['total_conf_reg'])})")
        if trip['total_transportation'] > 0:
            costs.append(f"ground transportation ({format_currency(trip['total_transportation'])})")

        cost_str = ""
        if costs:
            if len(costs) == 1:
                cost_str = f"Costs include {costs[0]}. "
            else:
                cost_str = f"Costs include {', '.join(costs[:-1])} and {costs[-1]}. "

        # International travel is always assumed to be conference travel
        narrative = f"This trip to {dest_escaped} will support {travelers} traveler(s) attending an international professional conference for {days} day(s) and {nights} night(s). "
        narrative += f"{cost_str}"
        narrative += f"The total cost for this international conference trip is {format_currency(total_cost)} (see Table~\\ref{{tab:international_travel}}, Trip {trip_idx}). "
        narrative += "International conference attendance is essential for disseminating research findings to the global scientific community, networking with international collaborators, staying current with worldwide developments in the field, and receiving feedback from leading researchers. "
        narrative += highlight_todo("[TODO: Specify conference name and location, e.g., 'International Conference on Machine Learning (ICML) in Vienna, Austria']")

        return narrative

    def generate_travel_section(self):
        """Section E: Travel with separate domestic and international tables"""
        total = self.data.cumulative_data.get('total_travel', {}).get('total', 0)
        domestic_total = self.data.cumulative_data.get('domestic_travel', {}).get('total', 0)
        international_total = self.data.cumulative_data.get('international_travel', {}).get('total', 0)

        section = f"\\subsection*{{E. Travel—{format_currency(total)}}}\n"
        section += "Travel funds are requested to support presenting research at conferences, conducting field work, visiting collaborating institutions, and disseminating research findings to the scientific community and stakeholders.\n\n"

        # E1: Domestic Travel
        section += f"\\subsubsection*{{E1. Domestic Travel—{format_currency(domestic_total)}}}\n"

        if self.data.domestic_travel:
            # Add yearly breakdown
            dom_yearly = []
            for i in range(self.data.years):
                year_amt = self.data.cumulative_data.get('domestic_travel', {}).get(f'year{i+1}', 0)
                if year_amt > 0:
                    dom_yearly.append(f"Year {i+1}: {format_currency(year_amt)}")
            if dom_yearly:
                section += f"Domestic travel totaling {format_currency(domestic_total)} is allocated as follows: {'; '.join(dom_yearly)}.\n\n"

            # Generate detailed paragraph for each domestic trip
            for idx, trip in enumerate(self.data.domestic_travel, 1):
                dest = trip['destination']
                travelers = int(trip['travelers'])
                days = int(trip['days'])
                nights = int(trip['nights'])
                total_cost = trip['cumulative']

                section += f"\\textbf{{Trip \\#{idx}:}} "
                section += self.get_domestic_trip_narrative(dest, travelers, days, nights, trip, total_cost, idx)
                section += "\n\n"

            # Add reference to table
            section += "See Table~\\ref{tab:domestic_travel} for a detailed breakdown of domestic travel costs.\n\n"
            section += self.generate_domestic_travel_table()
            section += "\n\n"
        else:
            section += "No domestic travel is budgeted for this project.\n\n"

        # E2: International Travel
        section += f"\\subsubsection*{{E2. International Travel—{format_currency(international_total)}}}\n"

        if self.data.international_travel:
            # Add yearly breakdown
            intl_yearly = []
            for i in range(self.data.years):
                year_amt = self.data.cumulative_data.get('international_travel', {}).get(f'year{i+1}', 0)
                if year_amt > 0:
                    intl_yearly.append(f"Year {i+1}: {format_currency(year_amt)}")
            if intl_yearly:
                section += f"International travel totaling {format_currency(international_total)} is allocated as follows: {'; '.join(intl_yearly)}.\n\n"

            # Generate detailed paragraph for each international trip (always conference)
            for idx, trip in enumerate(self.data.international_travel, 1):
                dest = trip['destination']
                travelers = int(trip['travelers'])
                days = int(trip['days'])
                nights = int(trip['nights'])
                total_cost = trip['cumulative']

                section += f"\\textbf{{Trip \\#{idx}:}} "
                section += self.get_international_trip_narrative(dest, travelers, days, nights, trip, total_cost, idx)
                section += "\n\n"

            # Add reference to table
            section += "See Table~\\ref{tab:international_travel} for a detailed breakdown of international travel costs.\n\n"
            section += self.generate_international_travel_table()
            section += "\n\n"
        else:
            section += "No international travel is budgeted for this project.\n\n"

        return section

    def generate_participant_costs_section(self):
        """Section F: Participant/Trainee Support Costs"""
        total = self.data.cumulative_data.get('participant_support', {}).get('total', 0)

        section = f"\\subsection*{{F. Participant Support Costs—"
        if total == 0:
            section += "N/A}\n"
            section += "No participant/trainee support costs are requested for this project.\n\n"
        else:
            section += f"{format_currency(total)}}}\n"
            section += highlight_todo("[TODO: Describe participant support costs including stipends, travel, subsistence]") + "\n\n"

        return section

    def generate_other_direct_costs_section(self):
        """Section G: Other Direct Costs - with subsections for each category"""
        total = self.data.cumulative_data.get('other_direct_costs', {}).get('total', 0)
        tuition = self.data.cumulative_data.get('tuition_remission', {}).get('total', 0)
        subaward_total = self.data.cumulative_data.get('total_subawards', {}).get('total', 0)

        # Calculate grand total including tuition and subawards
        grand_total = total + tuition + subaward_total

        section = f"\\subsection*{{G. Other Direct Costs—{format_currency(grand_total)}}}\n"

        if grand_total == 0:
            section += "No other direct costs are requested for this project.\n\n"
            return section

        # Get extracted line items from rows 177-192
        odc_items = getattr(self.data, 'odc_items', [])

        # Helper to format yearly breakdown with range grouping
        def format_yearly(item_yearly):
            year_data = []
            for i in range(self.data.years):
                amt = item_yearly.get(f'year{i+1}', 0)
                if amt > 0:
                    year_data.append((i+1, amt))
            return format_year_range(year_data) if year_data else None

        # Default justifications for standard category labels (lines 182-188)
        default_justifications = {
            'publication': "Publication costs are requested to support open-access publication fees and page charges for disseminating research findings in peer-reviewed journals.",
            'consultant': "Consultant services provide specialized expertise not available within the project team, ensuring high-quality guidance and evaluation of project activities.",
            'alteration': "Alterations and renovations are necessary to prepare research facilities for project-specific requirements and ensure compliance with safety and operational standards.",
            'computer': "Computer and software services support essential computational infrastructure, data storage, analysis tools, and cloud computing resources required for research activities.",
            'software': "Computer and software services support essential computational infrastructure, data storage, analysis tools, and cloud computing resources required for research activities.",
            'facility': "Facility rental and user fees provide access to specialized equipment, core facilities, and research infrastructure necessary for project objectives.",
            'user fee': "Facility rental and user fees provide access to specialized equipment, core facilities, and research infrastructure necessary for project objectives.",
            'human subject': "Participant compensation is essential for recruiting and retaining human subjects, ensuring adequate sample sizes, and recognizing participants' time and contribution to the research.",
        }

        # Group items by category for subsections
        if odc_items:
            # Organize items by category
            categories = {}
            for item in odc_items:
                category_raw = item.get('category', '').strip()
                if not category_raw:
                    continue
                if category_raw not in categories:
                    categories[category_raw] = []
                categories[category_raw].append(item)

            # Generate subsections for each category
            subsection_num = 1
            for category_raw, items in categories.items():
                category = escape_latex(category_raw)
                category_lower = category_raw.lower()

                # Calculate category total
                category_total = sum(item.get('total', 0) for item in items)

                # Create subsection header
                section += f"\\subsubsection*{{G{subsection_num}. {category}—{format_currency(category_total)}}}\n"

                # Add items under this category
                for item in items:
                    desc = escape_latex(item.get('description', '').strip())
                    item_total = item.get('total', 0)
                    yearly_str = format_yearly(item.get('yearly', {}))

                    # Build item description
                    item_text = f"{desc}, totalling {format_currency(item_total)}"
                    if yearly_str:
                        item_text += f" with a yearly breakdown of: {yearly_str}"
                    item_text += "."

                    section += item_text + " "

                # Add justification based on category
                is_materials = 'material' in category_lower and 'suppl' in category_lower

                if is_materials:
                    section += highlight_todo("[TODO: Justify need for these materials and supplies]")
                else:
                    # Check if category matches a standard label
                    matched_justification = None
                    for key, justification in default_justifications.items():
                        if key in category_lower:
                            matched_justification = justification
                            break

                    if matched_justification:
                        section += matched_justification
                    else:
                        section += highlight_todo("[TODO: Justify this expense]")

                section += "\n\n"
                subsection_num += 1
        else:
            # Fallback if no items extracted - show category totals with TODO
            subsection_num = 1
            if materials > 0:
                section += f"Materials and supplies totaling {format_currency(materials)} are requested. "
            if consultants > 0:
                section += f"Consultant services totaling {format_currency(consultants)} are requested. "
            if other > 0:
                section += f"Other expenses totaling {format_currency(other)} are requested. "
            section += highlight_todo("[TODO: Describe specific items, costs, and justifications]")
            section += "\n\n"

        # GX. Tuition Remission (numbered after ODC categories)
        if tuition > 0:
            section += f"\\subsubsection*{{G{subsection_num}. Tuition Remission—{format_currency(tuition)}}}\n"
            subsection_num += 1
            section += "Tuition is included as a graduate student benefit and is charged to projects in proportion to the amount of effort each graduate student will work on the project. "
            section += "Tuition charges are exempt from Facilities and Administrative (F\\&A) costs. "
            section += f"The tuition charge for graduate students is based upon the currently approved University tuition and fee schedules of \\$12,936 per student in Year 1 and includes an expected 5\\% annual increase. "
            section += f"Total tuition remission is calculated at {format_currency(tuition)} for all years.\n\n"

        # GX. Subawards (numbered after Tuition)
        if subaward_total > 0:
            section += f"\\subsubsection*{{G{subsection_num}. Subawards—{format_currency(subaward_total)}}}\n"

            # Get subaward organization names from extraction
            subaward_names = getattr(self.data, 'subaward_names', [])
            if subaward_names:
                org_name = escape_latex(subaward_names[0])
            else:
                org_name = highlight_todo("[TODO: Institution Name]")

            # Add year-by-year breakdown
            year_data = []
            for i in range(self.data.years):
                year_amt = self.data.cumulative_data.get('total_subawards', {}).get(f'year{i+1}', 0)
                if year_amt > 0:
                    year_data.append((i+1, year_amt))

            if year_data:
                section += f"Total subaward costs over {self.data.years} years are {format_currency(subaward_total)}, allocated as follows: {format_year_range(year_data)}. "

            section += f"This subaward supports critical project activities that will be conducted by collaborating investigators at {org_name}. "
            section += "The subaward institution will provide essential expertise, resources, and capabilities that complement the work performed at Northern Arizona University. "
            section += "Key personnel at the subaward institution include " + highlight_todo("[TODO: Name collaborating PI/Co-I and their role]") + ". "
            section += "\n\n"

            section += f"A detailed budget justification for {org_name} is provided in their detailed budget. "
            section += "The subaward includes support for personnel, supplies, equipment, and other costs necessary to complete their scope of work. "
            section += "Note: The first \\$25,000 per year of the subaward is subject to Northern Arizona University's indirect cost rate; subaward amounts exceeding \\$25,000 per year are exempt from indirect costs per federal regulations.\n\n"

        return section

    def generate_total_direct_costs_section(self):
        """Section H: Total Direct Costs"""
        direct = self.data.cumulative_data.get('total_direct_costs', {}).get('total', 0)

        section = f"\\subsection*{{H. Total Direct Costs—{format_currency(direct)}}}\n"
        section += f"The total direct costs for this project are {format_currency(direct)}, representing the sum of Sections A through G.\n\n"

        return section

    def generate_subaward_section_old(self):
        """Section H: Subawards"""
        total = self.data.cumulative_data.get('total_subawards', {}).get('total', 0)

        section = f"\\subsection*{{H. Subaward—"
        if total == 0:
            section += "N/A}\n"
            section += "No subawards are requested for this project.\n\n"
        else:
            section += f"{format_currency(total)}}}\n"

            # Get subaward organization names from extraction
            subaward_names = getattr(self.data, 'subaward_names', [])
            if subaward_names:
                org_name = escape_latex(subaward_names[0])  # Use first extracted name
            else:
                org_name = highlight_todo("[TODO: Institution Name]")

            # Add year-by-year breakdown
            year_data = []
            for i in range(self.data.years):
                year_amt = self.data.cumulative_data.get('total_subawards', {}).get(f'year{i+1}', 0)
                if year_amt > 0:
                    year_data.append((i+1, year_amt))

            if year_data:
                section += f"Total subaward costs over {self.data.years} years are {format_currency(total)}, allocated as follows: {format_year_range(year_data)}. "

            section += f"This subaward supports critical project activities that will be conducted by collaborating investigators at {org_name}. "
            section += "The subaward institution will provide essential expertise, resources, and capabilities that complement the work performed at Northern Arizona University. "
            section += "Key personnel at the subaward institution include " + highlight_todo("[TODO: Name collaborating PI/Co-I and their role]") + ". "
            section += "\n\n"

            section += f"A detailed budget justification for {org_name} is provided in their detailed budget. "
            section += "The subaward includes support for personnel, supplies, equipment, and other costs necessary to complete their scope of work. "
            section += "Note: The first \\$25,000 per year of the subaward is subject to Northern Arizona University's indirect cost rate; subaward amounts exceeding \\$25,000 per year are exempt from indirect costs per federal regulations.\n\n"

        return section

    def generate_indirect_section(self):
        """Section I: Indirect Costs with exact MTDC verbiage"""
        total = self.data.cumulative_data.get('total_indirect_costs', {}).get('total', 0)
        mtdc_base = self.data.cumulative_data.get('indirect_base', {}).get('total', 0)

        # Calculate rate
        rate = 52.5  # Default NAU rate
        if mtdc_base and mtdc_base > 0:
            rate = round((total / mtdc_base) * 100, 1)

        section = f"\\subsection*{{I. Indirect Costs—{format_currency(total)}}}\n"
        section += get_indirect_costs_text(rate, mtdc_base, total) + "\n\n"

        return section

    def generate_totals_section(self):
        """Section J: Total Project Costs - cumulative total only"""
        direct = self.data.cumulative_data.get('total_direct_costs', {}).get('total', 0)
        indirect = self.data.cumulative_data.get('total_indirect_costs', {}).get('total', 0)
        total = direct + indirect

        section = f"\\subsection*{{J. Total Project Costs—{format_currency(total)}}}\n"
        section += f"The total project cost is {format_currency(total)}.\n\n"

        return section

    def generate_complete_tex(self):
        """Assemble all sections into complete LaTeX document content"""
        sections = [
            self.generate_header(),
            self.generate_personnel_section(),
            self.generate_other_personnel_section(),
            self.generate_fringe_section(),
            self.generate_equipment_section(),
            self.generate_travel_section(),
            self.generate_participant_costs_section(),
            self.generate_other_direct_costs_section(),  # Now includes Tuition and Subaward as subsections
            self.generate_total_direct_costs_section(),  # NEW Section H
            self.generate_indirect_section(),  # Section I
            self.generate_totals_section()  # Section J
        ]

        return '\n\n'.join(sections)

    def save_standalone(self, output_path):
        """Save standalone compilable version with full LaTeX preamble matching main.tex formatting"""
        content = f"""\\documentclass[10pt]{{report}}
\\usepackage[left=0.5in, right=0.5in, top=0.5in, bottom=0.5in]{{geometry}}

% Remove all links (not allowed in grants)
\\usepackage[draft]{{hyperref}}
\\renewcommand{{\\href}}[2]{{#2}}
\\renewcommand{{\\url}}[1]{{\\nolinkurl{{#1}}}}
\\PassOptionsToPackage{{draft}}{{hyperref}}

% Font setup - Arial 10pt
\\usepackage{{fontspec}}
\\setmainfont{{Arial}}

\\usepackage{{graphicx}}
\\usepackage{{color,url}}
\\usepackage{{multirow}}
\\usepackage[skip=6pt, indent=0pt, parfill=0pt]{{parskip}}
\\usepackage[table]{{xcolor}}
\\usepackage{{array}}
\\usepackage[font={{small,sf}},labelfont=bf]{{caption}}

% Line spacing
\\linespread{{1}}

\\usepackage{{titlesec}}

\\setcounter{{secnumdepth}}{{3}}
\\setlength{{\\parindent}}{{0pt}}
\\counterwithout{{section}}{{chapter}}

% Chapter title with gray background shading
\\titleformat{{\\chapter}}[block]{{\\bf}}{{\\rlap{{\\color{{gray!20}}\\rule[-0.25cm]{{\\linewidth}}{{0.7cm}}}}}}{{0.1em}}{{\\MakeUppercase}}

% Numbering format
\\renewcommand\\thesection{{\\Alph{{section}}.}}
\\renewcommand\\thesubsection{{\\Alph{{section}}\\arabic{{subsection}}.}}
\\renewcommand\\thesubsubsection{{\\thesubsection\\arabic{{subsubsection}}}}

\\titleformat{{\\section}}{{\\normalfont\\bfseries}}{{\\thesection~}}{{0em}}{{\\MakeUppercase}}
\\titleformat{{\\subsection}}[runin]{{\\normalfont\\bfseries}}{{\\thesubsection~}}{{0em}}{{}}[:~ ]
\\titleformat{{\\subsubsection}}[runin]{{\\normalfont\\itshape}}{{\\underline\\thesubsubsection~}}{{0em}}{{}}[:~ ]
\\titleformat{{\\paragraph}}[runin]{{\\normalfont\\itshape}}{{}}{{0em}}{{}}[:~ ]

% Spacing
\\titlespacing{{\\chapter}}{{0pt}}{{-16pt}}{{0pt}}
\\titlespacing{{\\section}}{{0pt}}{{6pt}}{{0pt}}
\\titlespacing{{\\subsection}}{{0pt}}{{6pt}}{{0pt}}
\\titlespacing{{\\subsubsection}}{{0pt}}{{6pt}}{{0pt}}
\\titlespacing{{\\paragraph}}{{0pt}}{{6pt}}{{0pt}}

\\usepackage[belowskip=0pt,aboveskip=0pt]{{caption}}
\\setlength{{\\intextsep}}{{10pt plus 2pt minus 2pt}}

% Custom commands
\\newcommand{{\\TODO}}[1]{{{{\\color{{red}}{{\\bf [TODO: #1]}}}}}}

\\begin{{document}}

\\pagenumbering{{gobble}}
\\noindent

% Generated from {os.path.basename(self.data.filepath)}
% Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}
% Template: {self.data.years}-year project
% Script: generate_budget_justification.py

{self.generate_complete_tex()}

\\end{{document}}
"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

# ============================================================================
# MAIN FUNCTION
# ============================================================================

def main():
    """Main entry point for the script"""
    parser = argparse.ArgumentParser(
        description='Generate NIH Budget Justification (LaTeX and DOCX) from NAU Excel templates',
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog='''
Examples:
  python3 generate_budget_justification.py MyBudget.xlsx
  python3 generate_budget_justification.py MyBudget.xlsx -o output_directory
  python3 generate_budget_justification.py MyBudget.xlsx -v

Output Files:
  - MyBudget_BudgetJustification.tex (LaTeX source)
  - MyBudget_BudgetJustification.docx (Word document, requires pandoc)

Note:
  Install pandoc for DOCX generation: brew install pandoc (macOS) or apt-get install pandoc (Linux)
        '''
    )
    parser.add_argument('excel_file', help='Path to Excel budget file (.xlsx)')
    parser.add_argument('-o', '--output', default='.', help='Output directory (default: current directory)')
    parser.add_argument('-v', '--verbose', action='store_true', help='Verbose output showing extraction details')

    args = parser.parse_args()

    # Check if file exists
    if not os.path.isfile(args.excel_file):
        print(f"❌ Error: File '{args.excel_file}' not found.")
        return 1

    # Check if output directory exists
    if not os.path.isdir(args.output):
        print(f"Creating output directory: {args.output}")
        os.makedirs(args.output, exist_ok=True)

    print("\nAnalyzing budget template...")

    try:
        # Extract data from Excel
        extractor = BudgetExtractor(args.excel_file)
        extractor.extract_senior_personnel()
        extractor.extract_other_personnel()
        extractor.extract_travel()
        extractor.extract_cumulative()
        extractor.extract_subaward_names()
        extractor.extract_other_direct_costs_items()

        # Display extraction summary
        print(f"✓ Template type detected: {extractor.years}-year project")

        if args.verbose:
            # Senior personnel details
            sr_count = len(extractor.senior_personnel)
            if sr_count > 0:
                pi_count = sum(1 for p in extractor.senior_personnel if 'PI' in p['role'] and 'Co' not in p['role'])
                copi_count = sum(1 for p in extractor.senior_personnel if 'Co' in p['role'] and 'PI' in p['role'])
                senior_count = sr_count - pi_count - copi_count
                print(f"✓ Found {sr_count} senior personnel ({pi_count} PI, {copi_count} Co-PI, {senior_count} Senior Personnel)")
                for person in extractor.senior_personnel:
                    print(f"  - {person['name']} ({person['role']}): {person.get('pm_y1', 0)} PM/year")
            else:
                print("✓ Found 0 senior personnel")

            # Other personnel
            oth_count = len(extractor.other_personnel)
            print(f"✓ Found {oth_count} other personnel positions")
            for position in extractor.other_personnel:
                print(f"  - {position['role']}: {position['hours_week']} hrs/week")

            # Travel
            dom_count = len(extractor.domestic_travel)
            int_count = len(extractor.international_travel)
            print(f"✓ Found {dom_count} domestic trips, {int_count} international trips")
        else:
            print(f"✓ Found {len(extractor.senior_personnel)} senior personnel")
            print(f"✓ Found {len(extractor.other_personnel)} other personnel positions")
            print(f"✓ Found {len(extractor.domestic_travel)} domestic trips, {len(extractor.international_travel)} international trips")

        print(f"✓ Extracted cumulative data for {extractor.years} years")

        # Generate LaTeX file (self-contained standalone with full formatting)
        generator = LaTeXGenerator(extractor)

        basename = os.path.splitext(os.path.basename(args.excel_file))[0]
        standalone_path = os.path.join(args.output, f"{basename}_BudgetJustification.tex")

        generator.save_standalone(standalone_path)

        # Generate DOCX using Pandoc
        docx_path = os.path.join(args.output, f"{basename}_BudgetJustification.docx")
        try:
            # Check if pandoc is available
            result = subprocess.run(['pandoc', '--version'], capture_output=True, text=True)
            if result.returncode == 0:
                # Convert LaTeX to DOCX using Pandoc
                pandoc_cmd = [
                    'pandoc',
                    standalone_path,
                    '-f', 'latex',
                    '-t', 'docx',
                    '-o', docx_path,
                    '--standalone',
                    '--reference-doc=' + args.reference_docx if hasattr(args, 'reference_docx') and args.reference_docx else ''
                ]
                # Remove empty reference-doc argument if not provided
                pandoc_cmd = [arg for arg in pandoc_cmd if arg]

                subprocess.run(pandoc_cmd, check=True, capture_output=True)

                # Format the DOCX file to match PDF formatting
                if format_docx_file(docx_path):
                    docx_generated = True
                else:
                    docx_generated = False
                    if args.verbose:
                        print(f"\n⚠ Warning: DOCX file created but formatting failed.")
            else:
                docx_generated = False
        except (subprocess.CalledProcessError, FileNotFoundError) as e:
            docx_generated = False
            if args.verbose:
                print(f"\n⚠ Warning: Could not generate DOCX file. Pandoc may not be installed.")
                print(f"  Install with: brew install pandoc (macOS) or apt-get install pandoc (Linux)")

        print(f"\n✓ Generated files:")
        print(f"  - {standalone_path}")
        if docx_generated:
            print(f"  - {docx_path}")

        print(f"\nTo compile PDF (run twice to resolve cross-references):")
        print(f"  xelatex {os.path.basename(standalone_path)}")
        print(f"  xelatex {os.path.basename(standalone_path)}")

    except Exception as e:
        print(f"\n❌ Error processing file: {e}")
        if args.verbose:
            import traceback
            traceback.print_exc()
        return 1

    return 0

if __name__ == '__main__':
    sys.exit(main())
